import asyncio
import os
import json
from datetime import datetime
from core.utils.engine import engine

# 📄 [OFFICE-INTEGRATION]: Triệu hồi thư viện chuyên dụng
try:
    from docx import Document
    from docx.shared import Pt
    WORD_READY = True
except ImportError:
    WORD_READY = False

class StrategicWriter:
    """
    💎 [STRATEGIC-WRITER v2.0]: Đại Biên tập Chiến lược.
    Chuyên trách: Viết báo cáo, thuyết minh dự án và xuất bản tệp DOCX chuẩn Elite.
    """
    def __init__(self):
        self.output_dir = "intelligence/outputs/reports"
        os.makedirs(self.output_dir, exist_ok=True)

    async def write_report(self, **kwargs):
        """
        ✍️ [MISSION-EXECUTION]: Soạn thảo và xuất bản văn bản.
        """
        topic = kwargs.get("topic", "Báo cáo Dự án")
        context = kwargs.get("context", "Không có dữ liệu bổ sung")
        structure = kwargs.get("structure", "Tiêu chuẩn thuyết minh dự án")
        
        # 🧠 [REASONING]: Khởi tạo tư duy sáng tác
        prompt = f"""Bạn là BIÊN TẬP VIÊN CHIẾN LƯỢC cao cấp của JKAI Zenith.
Nhiệm vụ: Soạn thảo bản THUYẾT MINH DỰ ÁN chuyên sâu cho Master.

CHỦ ĐỀ: "{topic}"
BỐI CẢNH DỮ LIỆU:
{context}

YÊU CẦU:
1. Văn phong: Uy nghiêm, hào sảng, chuyên nghiệp (Văn phong Chính phủ/Tập đoàn lớn).
2. Cấu trúc: {structure} (Bao gồm: Đặt vấn đề, Phân tích kỹ thuật, Dự toán ngân sách, Kết luận).
3. MECE: Đảm bảo không chồng chéo, không bỏ sót.

Hãy viết nội dung báo cáo hoàn chỉnh!"""

        report_content = await engine.call_chat(
            messages=[{"role": "user", "content": prompt}],
            role="PLANNER",
            task_id=kwargs.get("task_id", "writing_task")
        )

        # 📄 [DOCX-EXPORT]: Xuất bản ra tệp Word
        file_path = "N/A"
        if WORD_READY:
            try:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"Bao_cao_{timestamp}.docx"
                file_path = os.path.join(self.output_dir, filename)
                
                doc = Document()
                doc.add_heading(topic, 0)
                
                # Định dạng nội dung
                for section in report_content.split("\n\n"):
                    p = doc.add_paragraph(section)
                    p.style.font.size = Pt(12)
                
                doc.save(file_path)
                msg = f"✅ [SUCCESS]: Báo cáo đã được xuất bản tại {file_path}."
            except Exception as e:
                msg = f"⚠️ [EXPORT-ERR]: Không thể tạo tệp Word: {e}. Trả về bản thô."
        else:
            msg = "⚠️ [MISSING-LIB]: Thiếu thư viện python-docx. Trả về bản thô."

        return {
            "status": "success",
            "content": report_content,
            "file_path": file_path,
            "msg": msg
        }

# Đăng ký instance
_instance = StrategicWriter()
write_report = _instance.write_report
