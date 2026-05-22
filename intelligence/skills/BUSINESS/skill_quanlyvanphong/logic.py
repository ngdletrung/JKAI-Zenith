import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import time
from core.config import settings

# =================================================================
# 📊 JKAI ZENITH: LOGIC QUẢN TRỊ VĂN PHÒNG (OFFICE SUITE)
# =================================================================

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def soan_thao_word(title, content, output_path=None):
    """
    Kiến tạo văn bản Word chuẩn ELITE STYLIST.
    """
    if not output_path:
        output_path = os.path.join(settings.OUTPUT_DIR, f"[ZENITH]_Elite_Doc_{int(time.time())}.docx")
    
    print(f"📄 [JKAI-OFFICE] Đang kiến tạo văn bản Đẳng cấp: {title}")
    
    doc = Document()
    
    # 🎨 THIẾT LẬP PHONG CÁCH CHUẨN (ELITE STYLE)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # 1. TIÊU ĐỀ TRANG TRỌNG (COVER TITLE)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x57) # Xanh Navy Zenith
    
    doc.add_paragraph().add_run("_" * 40).font.color.rgb = RGBColor(0x1B, 0x3A, 0x57)
    doc.add_paragraph() # Spacing
    
    # 2. XỬ LÝ NỘI DUNG ĐA TẦNG
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        if line.startswith('# '):
            p = doc.add_heading(line.replace('# ', ''), level=1)
            # p.style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x57)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('- '):
            p = doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.line_spacing = 1.15 # Khoảng cách dòng thoáng đãng
    
    # 3. CHÂN TRANG ĐẲNG CẤP (ELITE FOOTER)
    doc.add_paragraph("\n\n")
    footer_line = doc.add_paragraph()
    footer_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_line.add_run("________________________________\n")
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f_run = footer.add_run("PHÁT HÀNH BỞI TẬP ĐOÀN JKAI ZENITH\n")
    f_run.bold = True
    f_run.font.size = Pt(9)
    f_run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x57)
    
    f_date = footer.add_run(f"Ngày lập: {time.strftime('%d/%m/%Y')}")
    f_date.font.size = Pt(8)
    f_date.font.italic = True
    
    # Đảm bảo thư mục tồn tại
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    
    return {"status": "success", "msg": f"Đã kiến tạo văn bản thành công tại: {output_path}", "path": output_path}

def xuat_bao_cao_excel(data, title="Báo cáo Zenith", output_path=None):
    """
    Xuất báo cáo số liệu ra Excel.
    Data: List of dicts hoặc JSON string.
    """
    if not output_path:
        output_path = os.path.join(settings.OUTPUT_DIR, f"[ZENITH]_Bao_Cao_{int(time.time())}.xlsx")
        
    print(f"📊 [JKAI-OFFICE] Đang trích xuất báo cáo Excel: {title}")
    
    try:
        if isinstance(data, str):
            data = json.loads(data)
            
        df = pd.DataFrame(data)
        
        # Đảm bảo thư mục tồn tại
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Xuất Excel với định dạng chuyên nghiệp
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Data')
            
        return {"status": "success", "msg": f"Đã trích xuất báo cáo Excel thành công tại: {output_path}", "path": output_path}
    except Exception as e:
        return {"status": "error", "msg": f"Lỗi trích xuất Excel: {str(e)}"}

def doc_du_lieu_van_phong(file_path):
    """Đọc và trích xuất dữ liệu từ tệp văn phòng."""
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == ".docx":
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            return {"status": "success", "type": "Word", "content": text[:5000]}
        elif ext in [".xlsx", ".xls", ".csv"]:
            df = pd.read_excel(file_path) if ext != ".csv" else pd.read_csv(file_path)
            return {"status": "success", "type": "Excel", "content": df.head(20).to_dict()}
        else:
            return {"status": "error", "msg": "Định dạng tệp không được hỗ trợ."}
    except Exception as e:
        return {"status": "error", "msg": f"Lỗi đọc tệp: {str(e)}"}
