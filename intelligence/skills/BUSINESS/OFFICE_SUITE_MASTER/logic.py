import os
import time
import json
import logging
from typing import Dict, Any, List, Optional, Union
import pandas as pd
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import copy
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.chart import BarChart, LineChart, PieChart, Reference

# 🛰️ JKAI CORE IMPORTS
from core.utils.engine import engine
from core.utils.converter import converter
from core.utils import path_manager

logger = logging.getLogger("ZenithOfficeMaster")

NAVY = RGBColor(0x1B, 0x3A, 0x57)
GRAY = RGBColor(0x80, 0x80, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _require(payload: dict, *keys: str) -> Optional[str]:
    """Trả về câu hỏi bổ sung nếu thiếu thông tin bắt buộc."""
    missing = [k for k in keys if not payload.get(k)]
    if missing:
        labels = {
            "title": "tiêu đề văn bản",
            "content": "nội dung văn bản",
            "data": "dữ liệu bảng",
            "filename": "tên file xuất ra",
            "format": "định dạng file (docx/xlsx/pdf)",
            "columns": "danh sách cột",
            "rows": "danh sách dòng dữ liệu",
        }
        needed = ", ".join(labels.get(k, k) for k in missing)
        return f"Thiếu thông tin cần thiết: {needed}. Vui lòng cung cấp để tôi hoàn tất."
    return None


class ZenithOfficeMaster:
    def __init__(self):
        self.output_dir = self._resolve_output_dir()
        os.makedirs(self.output_dir, exist_ok=True)

    def _resolve_output_dir(self) -> str:
        """Resolve thư mục xuất kết quả về tọa độ tuyệt đối dùng chung mọi container."""
        candidates = []
        # Container JKAI: host root D:\Docker\JKAI được mount tại /workspace
        if os.path.isdir("/workspace"):
            candidates.append(os.path.join("/workspace", "workspace", "outputs"))
        # Host / container khác: theo path_manager (FILES_OUTPUT trong path_rules.md)
        try:
            pm = path_manager.get("FILES_OUTPUT")
            if pm:
                candidates.append(os.path.abspath(pm))
        except Exception:
            pass
        candidates.append(os.path.abspath("files/Output"))
        for cand in candidates:
            try:
                os.makedirs(cand, exist_ok=True)
                test_file = os.path.join(cand, ".jkk_w")
                with open(test_file, "w") as f:
                    f.write("ok")
                os.remove(test_file)
                return cand
            except Exception:
                continue
        return candidates[-1]

    def _get_path(self, filename: str, ext: str) -> str:
        clean_name = "".join([c if c.isalnum() or c in "._- " else "_" for c in str(filename)])
        return os.path.join(self.output_dir, f"{clean_name}.{ext}")

    # ─────────────────────────────────────────────────────────────
    # WORD (python-docx)
    # ─────────────────────────────────────────────────────────────

    def write_word(
        self,
        title: str = "",
        content: str = "",
        filename: str = "Zenith_Doc",
        cover: bool = False,
        author: str = "",
        toc: bool = False,
        **kwargs,
    ) -> str:
        """Kiến tạo văn bản Word chuẩn Elite. content hỗ trợ Markdown nhẹ (#, ##, ###, -, 1., |table|)."""
        if not title or not content:
            raise ValueError("Cần cung cấp cả 'title' và 'content' để tạo văn bản Word.")
        output_path = self._get_path(filename, "docx")
        doc = Document()

        # Margins chuẩn hành chính VN
        for section in doc.sections:
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        if cover:
            for _ in range(5):
                doc.add_paragraph()
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cp.add_run(title.upper())
            r.bold = True
            r.font.size = Pt(24)
            r.font.color.rgb = NAVY
            doc.add_paragraph()
            if author:
                ap = doc.add_paragraph()
                ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ar = ap.add_run(f"Tác giả: {author}")
                ar.font.size = Pt(14)
            dp = doc.add_paragraph()
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dr = dp.add_run(time.strftime('%d/%m/%Y'))
            dr.font.size = Pt(12)
            doc.add_page_break()
        else:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = NAVY
            doc.add_paragraph().add_run("_" * 50).font.color.rgb = NAVY

        self._parse_markdown_into_doc(doc, content)

        # Header: tên văn bản
        header = doc.sections[0].header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run(title[:40])
        hr.font.size = Pt(9)
        hr.font.color.rgb = GRAY

        # Footer: trang X
        footer = doc.sections[0].footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run("JKAI ZENITH SOVEREIGN | ").font.size = Pt(8)
        run = fp.add_run()
        run.font.size = Pt(8)
        fld = run._element
        instr = fld.makeelement(qn('w:fldSimple'), {qn('w:instr'): 'PAGE'})
        fld.append(instr)

        doc.save(output_path)
        return output_path

    def _parse_markdown_into_doc(self, doc: Document, content: str) -> None:
        """Chuyển Markdown đơn giản thành tài liệu Word có cấu trúc."""
        for line in content.split('\n'):
            line = line.rstrip()
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue
            if stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('|') and stripped.endswith('|'):
                self._parse_md_table(doc, stripped)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif stripped[0].isdigit() and '. ' in stripped[:4]:
                doc.add_paragraph(stripped, style='List Number')
            else:
                p = doc.add_paragraph(stripped)
                p.paragraph_format.line_spacing = 1.15

    def _parse_md_table(self, doc: Document, line: str) -> None:
        """Dựng bảng Word từ dòng markdown |a|b|c|."""
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows = [cells]
        # đoạn này được gọi từng dòng; để gọn, tạo bảng tối đa 1 dòng + tự động thêm ở caller nếu cần
        table = doc.add_table(rows=1, cols=len(cells))
        table.style = 'Light Shading Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, c in enumerate(cells):
            cell = table.rows[0].cells[i]
            cell.text = c
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

    def edit_word(self, file_path: str, replace: List[Dict] = None, append: str = "") -> str:
        """Chỉnh sửa văn bản Word có sẵn: thay thế chuỗi (giữ nguyên định dạng) hoặc thêm nội dung."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Không tìm thấy file: {file_path}")
        doc = Document(file_path)
        for rule in replace or []:
            old = rule.get("old", "")
            new = rule.get("new", "")
            if not old:
                continue
            for p in doc.paragraphs:
                if old in p.text:
                    for r in p.runs:
                        if old in r.text:
                            r.text = r.text.replace(old, new)
                            old = old.replace(old, new)
                            break
                    # xử lý chuỗi nằm rải rác nhiều run
                    if old and old in p.text:
                        p.text = p.text.replace(old, new)
        if append:
            doc.add_paragraph()
            doc.add_paragraph(append)
        doc.save(file_path)
        return file_path

    # ─────────────────────────────────────────────────────────────
    # EXCEL (openpyxl)
    # ─────────────────────────────────────────────────────────────

    def write_excel(
        self,
        data: Union[List[Dict], List[List]] = None,
        sheet_name: str = "Data",
        filename: str = "Zenith_Data",
        columns: List[str] = None,
        rows: List[List] = None,
        formatted: bool = True,
        **kwargs,
    ) -> str:
        """Kiến tạo bảng tính Excel chuyên nghiệp với header đẹp, số format, autofilter."""
        if data is None:
            data = self._build_table(columns, rows)
        if not data:
            raise ValueError("Cần cung cấp 'data' (danh sách dict) hoặc 'columns'+'rows'.")
        output_path = self._get_path(filename, "xlsx")

        if data and isinstance(data[0], dict):
            df = pd.DataFrame(data)
        else:
            df = pd.DataFrame(data[1:], columns=data[0] if isinstance(data[0], list) else None)

        wb = Workbook()
        ws = wb.active
        ws.title = (sheet_name or "Data")[:31]

        if formatted:
            header_fill = PatternFill(start_color="1B3A57", end_color="1B3A57", fill_type="solid")
            header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
            thin = Side(style="thin", color="B0B0B0")
            border = Border(left=thin, right=thin, top=thin, bottom=thin)

            for c_idx, col in enumerate(df.columns, start=1):
                cell = ws.cell(row=1, column=c_idx, value=str(col))
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = border

            for r_idx, row in enumerate(df.itertuples(index=False), start=2):
                for c_idx, val in enumerate(row, start=1):
                    cell = ws.cell(row=r_idx, column=c_idx, value=val)
                    cell.border = border
                    if isinstance(val, (int, float)):
                        cell.number_format = "#,##0"
                        cell.alignment = Alignment(horizontal="right")
                    else:
                        cell.alignment = Alignment(vertical="center")

            for c_idx, col in enumerate(df.columns, start=1):
                max_len = max(
                    [len(str(col))]
                    + [len(str(v)) if v is not None else 0 for v in df[col].tolist()]
                )
                ws.column_dimensions[get_column_letter(c_idx)].width = min(max(max_len + 2, 10), 40)

            ws.auto_filter.ref = ws.dimensions
            ws.freeze_panes = "A2"
        else:
            for r_idx, row in enumerate(df.itertuples(index=False), start=1):
                for c_idx, val in enumerate(row, start=1):
                    ws.cell(row=r_idx, column=c_idx, value=val)

        wb.save(output_path)
        return output_path

    def _build_table(self, columns: List[str], rows: List[List]) -> List[dict]:
        if not columns:
            raise ValueError("Cần cung cấp 'columns' (danh sách cột) khi không có 'data'.")
        out = []
        for row in rows or []:
            out.append({columns[i]: row[i] if i < len(row) else "" for i in range(len(columns))})
        return out

    def add_chart(
        self,
        data: List[Dict],
        chart_type: str = "bar",
        filename: str = "Zenith_Chart",
        title: str = "Biểu đồ",
        x_col: str = None,
        y_col: str = None,
        **kwargs,
    ) -> str:
        """Thêm biểu đồ bar/line/pie từ dữ liệu dạng dict."""
        if not data:
            raise ValueError("Cần cung cấp 'data' để vẽ biểu đồ.")
        output_path = self._get_path(filename, "xlsx")
        df = pd.DataFrame(data)
        wb = Workbook()
        ws = wb.active
        ws.title = "Data"
        for c_idx, col in enumerate(df.columns, start=1):
            ws.cell(row=1, column=c_idx, value=str(col))
        for r_idx, row in enumerate(df.itertuples(index=False), start=2):
            for c_idx, val in enumerate(row, start=1):
                ws.cell(row=r_idx, column=c_idx, value=val)

        x_axis = x_col or df.columns[0]
        y_axis = y_col or df.columns[-1]
        cat_ref = Reference(ws, min_col=df.columns.get_loc(x_axis) + 1, min_row=2,
                            max_row=ws.max_row)
        data_ref = Reference(ws, min_col=df.columns.get_loc(y_axis) + 1, min_row=1,
                             max_row=ws.max_row)

        if chart_type == "pie":
            chart = PieChart()
            chart.add_data(data_ref, titles_from_data=True)
            chart.set_categories(cat_ref)
        elif chart_type == "line":
            chart = LineChart()
            chart.add_data(data_ref, titles_from_data=True)
            chart.set_categories(cat_ref)
        else:
            chart = BarChart()
            chart.add_data(data_ref, titles_from_data=True)
            chart.set_categories(cat_ref)
        chart.title = title
        chart.height = 12
        chart.width = 24
        ws.add_chart(chart, "E2")
        wb.save(output_path)
        return output_path

    def edit_excel(self, file_path: str, cell: str = None, value=None, append_rows: List[list] = None) -> str:
        """Chỉnh sửa file Excel có sẵn: đặt giá trị ô hoặc thêm dòng."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Không tìm thấy file: {file_path}")
        wb = load_workbook(file_path)
        ws = wb.active
        if cell is not None:
            ws[cell] = value
        for row in append_rows or []:
            ws.append(row)
        wb.save(file_path)
        return file_path

    # ─────────────────────────────────────────────────────────────
    # PDF (reportlab)
    # ─────────────────────────────────────────────────────────────

    def write_pdf(
        self,
        title: str = "",
        content: str = "",
        filename: str = "Zenith_Report",
        author: str = "",
        **kwargs,
    ) -> str:
        """Kiến tạo PDF chuẩn A4 hỗ trợ tiếng Việt."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # Đăng ký font hỗ trợ Unicode cho tiếng Việt
        font_name = "Helvetica"
        for candidate in [
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/times.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ]:
            if os.path.exists(candidate):
                try:
                    pdfmetrics.registerFont(TTFont("JKAIUni", candidate))
                    font_name = "JKAIUni"
                    break
                except Exception:
                    pass

        output_path = self._get_path(filename, "pdf")
        doc = SimpleDocTemplate(
            output_path, pagesize=A4,
            leftMargin=3 * cm, rightMargin=2 * cm,
            topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("JTitle", parent=styles["Title"], fontName=font_name,
                                     fontSize=20, textColor=colors.HexColor("#1B3A57"))
        body_style = ParagraphStyle("JBody", parent=styles["BodyText"], fontName=font_name,
                                    fontSize=11, leading=15)
        head_style = ParagraphStyle("JHead", parent=styles["Heading1"], fontName=font_name,
                                    fontSize=14, textColor=colors.HexColor("#1B3A57"))

        story = [Paragraph(title, title_style), Spacer(1, 12)]
        if author:
            story.append(Paragraph(f"Tác giả: {author}", body_style))
            story.append(Spacer(1, 8))

        for line in content.split("\n"):
            s = line.strip()
            if not s:
                story.append(Spacer(1, 6))
            elif s.startswith("## "):
                story.append(Paragraph(s[3:], head_style))
            elif s.startswith("# "):
                story.append(Paragraph(s[2:], title_style))
            elif s.startswith("- "):
                story.append(Paragraph(f"• {s[2:]}", body_style))
            elif s.startswith("|"):
                self._append_pdf_table(story, s, body_style)
            else:
                story.append(Paragraph(s, body_style))

        doc.build(story)
        return output_path

    def _append_pdf_table(self, story, line: str, style) -> None:
        from reportlab.platypus import Table, TableStyle
        from reportlab.lib import colors
        cells = [c.strip() for c in line.strip("|").split("|")]
        t = Table([cells])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1B3A57")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(t)

    # ─────────────────────────────────────────────────────────────
    # READ + ORCHESTRATION
    # ─────────────────────────────────────────────────────────────

    async def read_any(self, file_path: str) -> str:
        try:
            return await converter.to_markdown(file_path)
        except Exception as e:
            return f"❌ [READ-ERR]: {str(e)}"

    def _missing_info_question(self, goal: str) -> Optional[str]:
        """Hỏi lại thông tin còn thiếu một cách thông minh dựa trên ý định."""
        g = goal.lower()
        need = []
        if any(k in g for k in ["hồ sơ", "ho so", "dossier", "bộ hồ sơ"]):
            need.append("loại hồ sơ (hồ sơ xin việc, dự thầu, thủ tục hành chính...)")
        if any(k in g for k in ["word", "docx", "văn bản", "van ban", "tờ trình", "to trinh", "công văn", "cong van", "báo cáo", "bao cao"]):
            need.append("tiêu đề + nội dung chính (đoạn mở đầu, các phần chính)")
        if any(k in g for k in ["excel", "xlsx", "bảng", "bang", "bảng tính", "bang tinh", "thống kê", "thong ke"]):
            need.append("danh sách cột + dữ liệu (hoặc nguồn dữ liệu)")
        if any(k in g for k in ["pdf", "file pdf", "xuất pdf", "xuat pdf"]):
            need.append("nội dung cần đưa vào PDF")
        if any(k in g for k in ["powerpoint", "slide", "trình bày", "trinh bay", "pptx", "thuyết trình", "thuyet trinh"]):
            need.append("số slide / các mục chính của bài trình bày")
        if need:
            return "Để tôi soạn đúng theo yêu cầu, xin vui lòng bổ sung: " + "; ".join(need) + "."
        return None

    async def process_office_mission(self, action: str, **kwargs) -> Dict[str, Any]:
        """Điều phối các tác vụ văn phòng."""
        goal = kwargs.get("goal") or kwargs.get("query") or ""

        # [EXTRACTED-PARAMS]: Hỗ trợ chuẩn execute_skill của pipeline (extracted_params là chuỗi mô tả)
        extracted = kwargs.get("extracted_params")
        if extracted and isinstance(extracted, str):
            try:
                parsed = json.loads(extracted)
                if isinstance(parsed, dict):
                    for k, v in parsed.items():
                        if k not in kwargs:
                            kwargs[k] = v
                    action = kwargs.get("action", action)
            except Exception:
                goal = goal or extracted
        if not action or action == "auto":
            if not goal:
                goal = kwargs.get("extracted_params", "") or goal
            action = _infer_action(kwargs) or "auto"

        # Thông minh: nếu chỉ có goal mơ hồ, hỏi lại thông tin thiếu
        if action in ("auto", "plan_dossier", "create_document") and not kwargs.get("content"):
            question = self._missing_info_question(goal)
            if question:
                return {"status": "need_info", "question": question, "documents": []}

        if action in ("read", "read_file"):
            content = await self.read_any(kwargs.get("file_path"))
            return {"status": "success", "content": content}

        elif action in ("write_word", "create_word", "docx"):
            missing = _require(kwargs, "title", "content")
            if missing:
                return {"status": "need_info", "question": missing}
            path = self.write_word(
                kwargs.get("title"), kwargs.get("content"),
                kwargs.get("filename", "Doc"),
                cover=kwargs.get("cover", False),
                author=kwargs.get("author", ""),
                toc=kwargs.get("toc", False),
            )
            return {"status": "success", "path": path, "format": "docx", "file": path}

        elif action in ("write_excel", "create_excel", "xlsx"):
            missing = _require(kwargs, "data")
            if not missing and not kwargs.get("data") and not (kwargs.get("columns") and kwargs.get("rows")):
                missing = "Thiếu dữ liệu bảng. Vui lòng cung cấp 'data' (danh sách dict) hoặc 'columns'+'rows'."
            if missing:
                return {"status": "need_info", "question": missing}
            path = self.write_excel(
                kwargs.get("data"),
                sheet_name=kwargs.get("sheet_name", "Data"),
                filename=kwargs.get("filename", "Data"),
                columns=kwargs.get("columns"),
                rows=kwargs.get("rows"),
                formatted=kwargs.get("formatted", True),
            )
            return {"status": "success", "path": path, "format": "xlsx", "file": path}

        elif action in ("write_pdf", "create_pdf", "pdf"):
            missing = _require(kwargs, "content")
            if missing:
                return {"status": "need_info", "question": missing}
            path = self.write_pdf(
                kwargs.get("title", kwargs.get("filename", "Report")),
                kwargs.get("content"),
                kwargs.get("filename", "Report"),
                author=kwargs.get("author", ""),
            )
            return {"status": "success", "path": path, "format": "pdf", "file": path}

        elif action in ("add_chart", "chart", "create_chart"):
            missing = _require(kwargs, "data")
            if missing:
                return {"status": "need_info", "question": missing}
            path = self.add_chart(
                kwargs.get("data"),
                chart_type=kwargs.get("chart_type", "bar"),
                filename=kwargs.get("filename", "Chart"),
                title=kwargs.get("title", "Biểu đồ"),
                x_col=kwargs.get("x_col"),
                y_col=kwargs.get("y_col"),
            )
            return {"status": "success", "path": path, "format": "xlsx", "file": path}

        elif action in ("edit_word", "edit"):
            path = self.edit_word(
                kwargs.get("file_path"),
                replace=kwargs.get("replace"),
                append=kwargs.get("append", ""),
            )
            return {"status": "success", "path": path}

        elif action in ("edit_excel",):
            path = self.edit_excel(
                kwargs.get("file_path"),
                cell=kwargs.get("cell"),
                value=kwargs.get("value"),
                append_rows=kwargs.get("append_rows"),
            )
            return {"status": "success", "path": path}

        return {"status": "error", "msg": f"Hành động '{action}' không được hỗ trợ. Các hành động: read, write_word, write_excel, write_pdf, add_chart, edit_word, edit_excel, create_document, plan_dossier."}


_instance = ZenithOfficeMaster()


def _infer_action(kwargs: dict) -> Optional[str]:
    """Tự suy ra action từ nội dung kwargs khi LLM không truyền 'action'."""
    if kwargs.get("action"):
        return kwargs["action"]
    if kwargs.get("method"):
        return kwargs["method"]
    if kwargs.get("file_path") or kwargs.get("path"):
        return "read"
    if kwargs.get("data") and (kwargs.get("chart_type") or kwargs.get("title") and "chart" in str(kwargs.get("chart_type", ""))):
        return "add_chart"
    if kwargs.get("data"):
        return "write_excel"
    if kwargs.get("content"):
        fmt = str(kwargs.get("format", "")).lower()
        if fmt in ("pdf",) or "pdf" in str(kwargs.get("filename", "")).lower():
            return "write_pdf"
        return "write_word"
    if kwargs.get("columns") and kwargs.get("rows"):
        return "write_excel"
    return None


async def execute_office_task(**kwargs):
    action = _infer_action(kwargs)
    if not action:
        return await _instance.process_office_mission("auto", **kwargs)
    kwargs.pop("action", None)
    kwargs.pop("method", None)
    return await _instance.process_office_mission(action, **kwargs)


async def OFFICE_SUITE_MASTER(**kwargs):
    """Entrypoint trùng tên skill_id để router khớp khi LLM gọi Action: OFFICE_SUITE_MASTER."""
    action = _infer_action(kwargs)
    if not action:
        return await _instance.process_office_mission("auto", **kwargs)
    kwargs.pop("action", None)
    kwargs.pop("method", None)
    return await _instance.process_office_mission(action, **kwargs)
